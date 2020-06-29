'''
autor: David A Herrera A
documentacion:
	este codigo prepara los archivos doc para procesar como texto
steps:
	remove hyperlink from file, accomplished from remhyp.py
	remove accents, accomplished from strip_accents(s)
	remove punctuations, accomplished with translator

'''

#import docx 	#to read docx files
#import nltk		#used for tokenizing
import re 
import string	#for replace
import unicodedata	#deal with accents
from ..clases.deHyperlink import get_docx_text

'''
inits
'''
translator = str.maketrans('', '', string.punctuation)	#removes punctuations

'''
functions
'''
def strip_hyperlink(path):
	return get_docx_text(path)
	
'''
def strip_accents(text):
	d = {"á":"a", "é":"e", "í":"i", "ó":"o", "ú":"u", "Á":"A", "É":"E", "Í":"I", "Ó":"O", "Ú":"U"}
	# Create a regular expression  from the dictionary keys
	regex = re.compile("(%s)" % "|".join(map(re.escape, d.keys())))
	
	# For each match, look-up corresponding value in dictionary
	return regex.sub(lambda mo: dict[mo.string[mo.start():mo.end()]], text) 
	
	
or ''.join(c for c in unidecode(mystring).lower() if ord(c) in range(97,123) or ord(c)==32).lstrip().rstrip()
'''
def strip_accents(str):
	return ''.join(c for c in unicodedata.normalize('NFD', str) if unicodedata.category(c) != 'Mn')

def strip_puncs(str):
	return str.translate(translator)

def get_text(path):
	return strip_accents(strip_puncs(strip_hyperlink(path)))

def strip_spaces(str):
	return ' '.join(str.split())
	
	
'''
import re
import unicodedata

def strip_accents(text):
    """
    Strip accents from input String.

    :param text: The input string.
    :type text: String.

    :returns: The processed String.
    :rtype: String.
    """
    try:
        text = unicode(text, 'utf-8')
    except NameError: # unicode is a default on python 3 
        pass
    text = unicodedata.normalize('NFD', text)
    text = text.encode('ascii', 'ignore')
    text = text.decode("utf-8")
    return str(text)

def text_to_id(text):
    """
    Convert input text to id.

    :param text: The input string.
    :type text: String.

    :returns: The processed String.
    :rtype: String.
    """
    text = strip_accents(text.lower())
    text = re.sub('[ ]+', '_', text)
    text = re.sub('[^0-9a-zA-Z_-]', '', text)
    return text
'''
'''
http://stackoverflow.com/questions/15261793/python-efficient-method-to-replace-accents-%C3%A9-to-e-remove-a-za-z-d-s-and
import collections
import string

table = collections.defaultdict(lambda: None)
table.update({
    ord('é'):'e',
    ord('ô'):'o',
    ord(' '):' ',
    ord('\N{NO-BREAK SPACE}'): ' ',
    ord('\N{EN SPACE}'): ' ',
    ord('\N{EM SPACE}'): ' ',
    ord('\N{THREE-PER-EM SPACE}'): ' ',
    ord('\N{FOUR-PER-EM SPACE}'): ' ',
    ord('\N{SIX-PER-EM SPACE}'): ' ',
    ord('\N{FIGURE SPACE}'): ' ',
    ord('\N{PUNCTUATION SPACE}'): ' ',
    ord('\N{THIN SPACE}'): ' ',
    ord('\N{HAIR SPACE}'): ' ',
    ord('\N{ZERO WIDTH SPACE}'): ' ',
    ord('\N{NARROW NO-BREAK SPACE}'): ' ',
    ord('\N{MEDIUM MATHEMATICAL SPACE}'): ' ',
    ord('\N{IDEOGRAPHIC SPACE}'): ' ',
    ord('\N{IDEOGRAPHIC HALF FILL SPACE}'): ' ',
    ord('\N{ZERO WIDTH NO-BREAK SPACE}'): ' ',
    ord('\N{TAG SPACE}'): ' ',
    })
table.update(dict(zip(map(ord,string.ascii_uppercase), string.ascii_lowercase)))
table.update(dict(zip(map(ord,string.ascii_lowercase), string.ascii_lowercase)))
table.update(dict(zip(map(ord,string.digits), string.digits)))

print('123 fôé BAR҉'.translate(table,))

'''	

	
'''
fulltext = []
for para in doc.paragraphs:			#iterate thru all paragraphs read
	fulltext.append(para.text)
'\n'.join(fullText)
tokens = nltk.word_tokenize(fulltext)
'''


'''
fuentes de codigo
'''
		
	
'''
http://stackoverflow.com/questions/25228106/how-to-extract-text-from-an-existing-docx-file-using-python-docx

import docx

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)
'''


'''
http://stackoverflow.com/questions/15057945/how-do-i-tokenize-a-string-sentence-in-nltk

>>> import nltk
>>> sentence = """At eight o'clock on Thursday morning
... Arthur didn't feel very good."""
>>> tokens = nltk.word_tokenize(sentence)
>>> tokens
['At', 'eight', "o'clock", 'on', 'Thursday', 'morning',
'Arthur', 'did', "n't", 'feel', 'very', 'good', '.']
'''

'''
http://stackoverflow.com/questions/34293875/how-to-remove-punctuation-marks-from-a-string-in-python-3-x-using-translate

	
You have to create a translation table using maketrans that you pass to the str.translate method.

In Python 3.1 and newer, maketrans is now a static-method on the str type, so you can use it to create a translation of each punctuation you want to None.

import string

# Thanks to Martijn Pieters for this improved version

# This uses the 3-argument version of str.maketrans
# with arguments (x, y, z) where 'x' and 'y'
# must be equal-length strings and characters in 'x'
# are replaced by characters in 'y'. 'z'
# is a string (string.punctuation here)
# where each character in the string is mapped
# to None
translator = str.maketrans('', '', string.punctuation)

# This is an alternative that creates a dictionary mapping
# of every character from string.punctuation to None (this will
# also work)
#translator = str.maketrans(dict.fromkeys(string.punctuation))

s = 'string with "punctuation" inside of it! Does this work? I hope so.'

# pass the translator to the string's translate method.
print(s.translate(translator))
'''


'''
#http://stackoverflow.com/questions/15175142/how-can-i-do-multiple-substitutions-using-regex-in-python
import re 

def multiple_replace(dict, text):
  # Create a regular expression  from the dictionary keys
  regex = re.compile("(%s)" % "|".join(map(re.escape, dict.keys())))

  # For each match, look-up corresponding value in dictionary
  return regex.sub(lambda mo: dict[mo.string[mo.start():mo.end()]], text) 

if __name__ == "__main__": 

  text = "Larry Wall is the creator of Perl"

  dict = {
    "Larry Wall" : "Guido van Rossum",
    "creator" : "Benevolent Dictator for Life",
    "Perl" : "Python",
  } 

  print multiple_replace(dict, text)
'''
